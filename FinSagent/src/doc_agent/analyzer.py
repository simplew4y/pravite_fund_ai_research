"""Upload-time document classification and summary agent.

This module is intentionally independent of the main AgenticRAG workflow. It is
small enough for prototype testing, but returns structured output that can later
be persisted, indexed, or routed into the MAS pipeline.
"""

from __future__ import annotations

import asyncio
import json
import logging
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


DOC_TYPES: Dict[str, str] = {
    "financial_report": "财报",
    "industry_report": "行业报告",
    "news": "新闻",
    "event_commentary": "事件点评",
    "company_research_valuation": "公司研究和估值报告",
    "meeting_minutes": "会议纪要",
    "other": "其他",
}


TYPE_KEYWORDS: Dict[str, List[str]] = {
    "financial_report": [
        "annual report",
        "quarterly report",
        "form 10-k",
        "form 10-q",
        "consolidated statements",
        "balance sheet",
        "income statement",
        "cash flow",
        "revenue",
        "gross profit",
        "net income",
        "年报",
        "季报",
        "财报",
        "财务报表",
        "资产负债表",
        "现金流量表",
        "利润表",
        "营业收入",
        "净利润",
    ],
    "industry_report": [
        "industry report",
        "market size",
        "market share",
        "cagr",
        "competitive landscape",
        "tam",
        "sam",
        "行业报告",
        "行业研究",
        "市场规模",
        "竞争格局",
        "产业链",
        "渗透率",
        "复合增长率",
    ],
    "news": [
        "press release",
        "news release",
        "reuters",
        "bloomberg",
        "announced",
        "reported",
        "新闻",
        "快讯",
        "发布",
        "公告称",
        "据悉",
        "今日",
    ],
    "event_commentary": [
        "event comment",
        "first take",
        "flash note",
        "quick take",
        "事件点评",
        "快评",
        "点评",
        "影响",
        "催化",
        "短评",
        "点评报告",
    ],
    "company_research_valuation": [
        "equity research",
        "target price",
        "valuation",
        "dcf",
        "ev/ebitda",
        "buy rating",
        "hold rating",
        "sell rating",
        "initiation",
        "公司研究",
        "深度报告",
        "估值",
        "目标价",
        "评级",
        "买入",
        "增持",
        "中性",
        "DCF",
        "市盈率",
        "EV/EBITDA",
    ],
    "meeting_minutes": [
        "meeting minutes",
        "minutes of meeting",
        "attendees",
        "q&a",
        "qa transcript",
        "会议纪要",
        "调研纪要",
        "纪要",
        "参会",
        "问答",
        "Q：",
        "A：",
        "Q:",
        "A:",
    ],
}


@dataclass
class ExtractedDocument:
    filename: str
    text: str
    extraction_method: str
    warnings: List[str]


class DocumentTriageAgent:
    """Classify an uploaded document and produce a compact structured summary."""

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self.max_chars = int(self.config.get("doc_agent_max_chars", 24000))
        self.summary_language = str(self.config.get("doc_agent_summary_language", "中文"))

    async def analyze_upload(
        self,
        filename: str,
        content: bytes,
        *,
        use_llm: bool = True,
    ) -> Dict[str, Any]:
        extracted = extract_document_text(filename, content)
        text = extracted.text.strip()
        heuristic = classify_by_keywords(filename, text)
        if not text:
            return self._empty_result(filename, extracted, heuristic)

        llm_result: Optional[Dict[str, Any]] = None
        if use_llm and self._can_call_llm():
            try:
                llm_result = await self._analyze_with_llm(filename, text, heuristic)
            except Exception as exc:
                logger.warning("Document triage LLM call failed: %s", exc, exc_info=True)
                extracted.warnings.append(f"LLM analysis failed; heuristic fallback used: {type(exc).__name__}: {exc}")

        if llm_result:
            result = normalize_llm_result(llm_result, heuristic)
            result["llm_used"] = True
        else:
            result = heuristic_summary(filename, text, heuristic)
            result["llm_used"] = False

        result.update(
            {
                "filename": filename,
                "text_chars": len(text),
                "extraction_method": extracted.extraction_method,
                "warnings": extracted.warnings,
            }
        )
        return result

    def _can_call_llm(self) -> bool:
        return bool(self.config.get("llm_base_url") and self.config.get("llm_model_name"))

    async def _analyze_with_llm(
        self,
        filename: str,
        text: str,
        heuristic: Dict[str, Any],
    ) -> Dict[str, Any]:
        try:
            from openai import AsyncOpenAI
        except Exception as exc:  # pragma: no cover - depends on runtime env
            raise RuntimeError("openai package is not available") from exc

        client = AsyncOpenAI(
            api_key=str(self.config.get("llm_api_key", "EMPTY")),
            base_url=self.config.get("llm_base_url"),
            timeout=float(self.config.get("doc_agent_llm_timeout_seconds", self.config.get("llm_timeout_seconds", 90))),
            max_retries=int(self.config.get("doc_agent_llm_max_retries", 1)),
        )
        prompt = build_triage_prompt(
            filename=filename,
            text=text[: self.max_chars],
            heuristic=heuristic,
            language=self.summary_language,
        )
        response = await client.chat.completions.create(
            model=str(self.config.get("llm_model_name")),
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=float(self.config.get("doc_agent_temperature", 0.1)),
            max_tokens=int(self.config.get("doc_agent_max_tokens", 1800)),
        )
        content = response.choices[0].message.content or "{}"
        return json.loads(_strip_code_fence(content))

    def _empty_result(
        self,
        filename: str,
        extracted: ExtractedDocument,
        heuristic: Dict[str, Any],
    ) -> Dict[str, Any]:
        return {
            "filename": filename,
            "doc_type": heuristic.get("doc_type", "other"),
            "doc_type_label": DOC_TYPES.get(heuristic.get("doc_type", "other"), "其他"),
            "confidence": "low",
            "summary": "未能从文件中抽取到可分析文本。",
            "summary_short": "文本抽取为空。",
            "key_points": [],
            "entities": [],
            "periods": [],
            "evidence_signals": heuristic.get("evidence_signals", []),
            "suggested_next_actions": ["检查文件是否为扫描件；如是扫描 PDF，需要先做 OCR。"],
            "text_chars": 0,
            "extraction_method": extracted.extraction_method,
            "warnings": extracted.warnings or ["No text extracted"],
            "llm_used": False,
        }


async def analyze_document_bytes(
    filename: str,
    content: bytes,
    config: Optional[Dict[str, Any]] = None,
    *,
    use_llm: bool = True,
) -> Dict[str, Any]:
    return await DocumentTriageAgent(config).analyze_upload(filename, content, use_llm=use_llm)


def extract_document_text(filename: str, content: bytes) -> ExtractedDocument:
    suffix = Path(filename or "").suffix.lower()
    warnings: List[str] = []
    if suffix == ".pdf":
        return _extract_pdf(filename, content)
    if suffix == ".docx":
        text = _extract_docx(content, warnings)
        return ExtractedDocument(filename, text, "python-docx", warnings)
    if suffix in {".txt", ".md", ".csv", ".json", ".html", ".htm", ".xml"}:
        text, encoding = _decode_bytes(content)
        if suffix in {".html", ".htm", ".xml"}:
            text = _strip_markup(text)
        return ExtractedDocument(filename, text, f"raw_text:{encoding}", warnings)

    text, encoding = _decode_bytes(content)
    if not text.strip():
        warnings.append(f"Unsupported file extension: {suffix or '(none)'}")
    return ExtractedDocument(filename, text, f"raw_bytes:{encoding}", warnings)


def _extract_pdf(filename: str, content: bytes) -> ExtractedDocument:
    warnings: List[str] = []
    text, method = _extract_pdf_with_pypdf(content, warnings)
    if not text.strip():
        text, method = _extract_pdf_with_pdfplumber(content, warnings)
    if not text.strip():
        text, method = _extract_pdf_with_fitz(content, warnings)
    return ExtractedDocument(filename, text, method, warnings)


def _extract_pdf_with_pypdf(content: bytes, warnings: List[str]) -> Tuple[str, str]:
    try:
        from io import BytesIO
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        pages = []
        for idx, page in enumerate(reader.pages):
            try:
                pages.append(page.extract_text() or "")
            except Exception as exc:
                warnings.append(f"pypdf page {idx + 1} failed: {exc}")
        return "\n\n".join(pages), "pypdf"
    except Exception as exc:
        warnings.append(f"pypdf unavailable or failed: {exc}")
        return "", "pypdf_failed"


def _extract_pdf_with_pdfplumber(content: bytes, warnings: List[str]) -> Tuple[str, str]:
    try:
        from io import BytesIO
        import pdfplumber

        pages = []
        with pdfplumber.open(BytesIO(content)) as pdf:
            for idx, page in enumerate(pdf.pages):
                try:
                    pages.append(page.extract_text() or "")
                except Exception as exc:
                    warnings.append(f"pdfplumber page {idx + 1} failed: {exc}")
        return "\n\n".join(pages), "pdfplumber"
    except Exception as exc:
        warnings.append(f"pdfplumber unavailable or failed: {exc}")
        return "", "pdfplumber_failed"


def _extract_pdf_with_fitz(content: bytes, warnings: List[str]) -> Tuple[str, str]:
    try:
        import fitz

        pages = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page in doc:
                pages.append(page.get_text("text") or "")
        return "\n\n".join(pages), "pymupdf"
    except Exception as exc:
        warnings.append(f"pymupdf unavailable or failed: {exc}")
        return "", "pymupdf_failed"


def _extract_docx(content: bytes, warnings: List[str]) -> str:
    try:
        from io import BytesIO
        from docx import Document

        doc = Document(BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        table_cells: List[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    table_cells.append(" | ".join(cells))
        return "\n".join(paragraphs + table_cells)
    except Exception as exc:
        warnings.append(f"python-docx unavailable or failed: {exc}")
        return ""


def _decode_bytes(content: bytes) -> Tuple[str, str]:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return content.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace"), "utf-8:replace"


def _strip_markup(text: str) -> str:
    text = re.sub(r"<script\b[^>]*>.*?</script>", " ", text, flags=re.I | re.S)
    text = re.sub(r"<style\b[^>]*>.*?</style>", " ", text, flags=re.I | re.S)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text)


def classify_by_keywords(filename: str, text: str) -> Dict[str, Any]:
    haystack = f"{filename}\n{text[:6000]}".lower()
    scores: Dict[str, int] = {}
    signals: Dict[str, List[str]] = {}
    for doc_type, keywords in TYPE_KEYWORDS.items():
        score = 0
        matched: List[str] = []
        for kw in keywords:
            kw_lower = kw.lower()
            count = haystack.count(kw_lower)
            if count:
                score += min(count, 5)
                matched.append(kw)
        scores[doc_type] = score
        signals[doc_type] = matched[:8]

    best_type = max(scores, key=scores.get) if scores else "other"
    best_score = scores.get(best_type, 0)
    if best_score <= 0:
        best_type = "other"
    sorted_scores = sorted(scores.items(), key=lambda item: item[1], reverse=True)
    second_score = sorted_scores[1][1] if len(sorted_scores) > 1 else 0
    if best_score >= 8 and best_score >= second_score * 1.5:
        confidence = "high"
    elif best_score >= 3:
        confidence = "medium"
    else:
        confidence = "low"

    return {
        "doc_type": best_type,
        "doc_type_label": DOC_TYPES.get(best_type, "其他"),
        "confidence": confidence,
        "scores": scores,
        "evidence_signals": signals.get(best_type, []),
    }


def heuristic_summary(filename: str, text: str, heuristic: Dict[str, Any]) -> Dict[str, Any]:
    clean = normalize_text(text)
    sentences = split_sentences(clean)
    key_points = pick_key_points(sentences)
    summary = " ".join(sentences[:4]).strip()
    if not summary:
        summary = clean[:500].strip()
    return {
        "doc_type": heuristic.get("doc_type", "other"),
        "doc_type_label": heuristic.get("doc_type_label", "其他"),
        "confidence": heuristic.get("confidence", "low"),
        "summary": summary[:1200],
        "summary_short": summary[:180],
        "key_points": key_points,
        "entities": extract_entities(filename, clean),
        "periods": extract_periods(clean),
        "evidence_signals": heuristic.get("evidence_signals", []),
        "suggested_next_actions": ["人工抽检分类结果；如需高质量摘要，请配置 LLM 后重新分析。"],
    }


def normalize_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def split_sentences(text: str) -> List[str]:
    raw = re.split(r"(?<=[。！？.!?])\s*|[\r\n]+", text)
    return [s.strip() for s in raw if len(s.strip()) >= 12][:40]


def pick_key_points(sentences: List[str]) -> List[str]:
    priority = re.compile(
        r"收入|营收|利润|毛利|现金流|市场|增长|同比|环比|估值|目标价|评级|风险|guidance|revenue|profit|margin|valuation|risk|growth",
        re.I,
    )
    picked = [s for s in sentences if priority.search(s)]
    if len(picked) < 3:
        picked.extend(s for s in sentences if s not in picked)
    return [s[:240] for s in picked[:5]]


def extract_entities(filename: str, text: str) -> List[str]:
    candidates: List[str] = []
    candidates.extend(re.findall(r"\b[A-Z]{2,6}\b", filename + " " + text[:3000]))
    candidates.extend(re.findall(r"[\u4e00-\u9fffA-Za-z0-9·&（）()]{2,30}(?:公司|集团|股份|汽车|科技|证券|银行)", text[:3000]))
    seen = set()
    entities = []
    for item in candidates:
        normalized = item.strip()
        if normalized and normalized not in seen:
            seen.add(normalized)
            entities.append(normalized)
    return entities[:12]


def extract_periods(text: str) -> List[str]:
    patterns = [
        r"20\d{2}年(?:第[一二三四1-4]季度|Q[1-4]|上半年|下半年|全年)?",
        r"20\d{2}\s?(?:Q[1-4]|H[12]|FY)?",
        r"(?:Q[1-4]|H[12]|FY)\s?20\d{2}",
    ]
    found: List[str] = []
    for pattern in patterns:
        found.extend(re.findall(pattern, text[:5000], flags=re.I))
    seen = set()
    periods = []
    for item in found:
        value = item.strip()
        if value and value not in seen:
            seen.add(value)
            periods.append(value)
    return periods[:10]


def build_triage_prompt(filename: str, text: str, heuristic: Dict[str, Any], language: str) -> str:
    allowed = "\n".join(f"- {key}: {label}" for key, label in DOC_TYPES.items())
    return f"""你是金融文档上传后的轻量分诊 agent。请根据文件名和正文片段，判断文档类型并生成摘要。

允许的 doc_type 只能是：
{allowed}

规则：
- 只输出 JSON，不要 Markdown。
- doc_type 必须从允许列表中选择一个。
- confidence 只能是 high、medium、low。
- summary 用{language}输出，要求短、准、覆盖文档主题、关键结论、重要数字/期间。
- 如果文本像扫描件或信息不足，要在 warnings 或 suggested_next_actions 中说明。
- 不要编造正文没有支持的数字或结论。

文件名：
{filename}

规则初判：
{json.dumps(heuristic, ensure_ascii=False)}

正文片段：
\"\"\"
{text}
\"\"\"

请输出如下 JSON schema：
{{
  "doc_type": "financial_report|industry_report|news|event_commentary|company_research_valuation|meeting_minutes|other",
  "doc_type_label": "中文类型名",
  "confidence": "high|medium|low",
  "summary_short": "一句话摘要，80字以内",
  "summary": "一段摘要，200-500字",
  "key_points": ["3-6条要点"],
  "entities": ["公司、行业、股票代码、人物等"],
  "periods": ["涉及期间，如2024年、2025Q1"],
  "evidence_signals": ["判断类型的文本信号"],
  "suggested_next_actions": ["后续处理建议"]
}}"""


def normalize_llm_result(payload: Dict[str, Any], heuristic: Dict[str, Any]) -> Dict[str, Any]:
    doc_type = str(payload.get("doc_type") or heuristic.get("doc_type") or "other")
    if doc_type not in DOC_TYPES:
        doc_type = heuristic.get("doc_type", "other")
    confidence = str(payload.get("confidence") or heuristic.get("confidence") or "low").lower()
    if confidence not in {"high", "medium", "low"}:
        confidence = "low"
    return {
        "doc_type": doc_type,
        "doc_type_label": str(payload.get("doc_type_label") or DOC_TYPES.get(doc_type, "其他")),
        "confidence": confidence,
        "summary_short": str(payload.get("summary_short") or "")[:300],
        "summary": str(payload.get("summary") or "")[:2400],
        "key_points": _coerce_str_list(payload.get("key_points"), limit=8),
        "entities": _coerce_str_list(payload.get("entities"), limit=16),
        "periods": _coerce_str_list(payload.get("periods"), limit=12),
        "evidence_signals": _coerce_str_list(payload.get("evidence_signals") or heuristic.get("evidence_signals"), limit=12),
        "suggested_next_actions": _coerce_str_list(payload.get("suggested_next_actions"), limit=8),
    }


def _coerce_str_list(value: Any, limit: int) -> List[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()][:limit]
    if isinstance(value, str) and value.strip():
        return [value.strip()][:limit]
    return []


def _strip_code_fence(content: str) -> str:
    content = content.strip()
    if content.startswith("```"):
        content = re.sub(r"^```(?:json)?\s*", "", content)
        content = re.sub(r"\s*```$", "", content)
    return content.strip()
